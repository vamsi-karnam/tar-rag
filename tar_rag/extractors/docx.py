"""DOCX text extractor using ``python-docx`` (optional dependency).

Install with ``pip install "tar-rag[docx]"``. Extracts paragraph text plus
table cell text, in document order.
"""

from __future__ import annotations

from ..errors import MissingExtractorDependency
from .base import TextExtractor


class DocxTextExtractor(TextExtractor):
    name = "DocxTextExtractor"

    def extract(self, file_path: str) -> str:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise MissingExtractorDependency(
                "DocxTextExtractor requires the 'python-docx' package. "
                "Install with: pip install \"tar-rag[docx]\""
            ) from exc

        document = Document(file_path)
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
